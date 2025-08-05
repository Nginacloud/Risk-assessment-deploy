import streamlit as st
import pandas as pd
from collections import defaultdict
import re
import pdfplumber
import docx 
import fitz  # PyMuPDF

st.set_page_config(page_title="Risk Assessment Tool", layout="centered")
st.title("Risk Assessment Tool")

# FILE TEXT EXTRACTOR 
def extract_text(file, password=None):
    if file.type == "application/pdf":
        try:
            file_bytes = file.read()
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            if doc.needs_pass and (not password or not doc.authenticate(password)):
                raise ValueError("PDF is encrypted and password is missing or incorrect.")
            #return "\n".join(page.get_text() for page in doc)
            extracted_text = "\n".join(page.get_text() for page in doc)
        except Exception as e:
            st.error(f"Failed to extract text from PDF: {e}")
            return ""
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(file)
        #return "\n".join(p.text for p in doc.paragraphs)
        extracted_text = "\n".join(p.text for p in doc.paragraphs)
    else:
        #return file.read().decode("utf-8")
        extracted_text = file.read().decode("utf-8")

    name_match = re.search(r"REPORTED NAMES:\s*(.*)", extracted_text)
    customer_name = name_match.group(1).strip() if name_match else None
    return extracted_text #customer_name

# M-PESA SECTION 
import re

def categorize_mpesa(details):
    text = details.lower().strip()

    # Define category patterns
    fuel_keywords = r"(fuel|petroleum|gas|diesel|oil|petrol|shell|totalenergies|petrol|rubis|Ola Energy|Energies|Kobil|KenolKobil|Astrol|Lake oil)"
    shopping_keywords = r"(supermarket|quickmart|naivas|chandarana|kaluu foods|nguku wholesalers|Clean Shelf|Magunas|tuskys|carrefour)"
    utility_keywords = r"(kplc|electric|prepaid|expressway|water)"
    airtime_keywords = r"\bairtime\b|\bbundle\b"
    betting_keywords = r"(betika|sportpesa|odibet|jackpot)"
    paybill_keywords = r"(pay bill|paybill)"
    buy_goods_keywords = r"(buy goods|merchant payment|till)"
    agent_keywords = r"(withdraw|agent)"
    salary_keywords = r"\bpayment from\b"#r"(salary payment|business payment|equity bulk|kopo kopo)"
    loan_keywords = r"(mpesa overdraw|od loan repayment)"
    WatuCredit_keywords = r"\bwatu credit\b"
    PlatinumCredit_keywords = r"\bplatinum\b"
    Mogocredit_keywords = r"\bmogo\b"
    MomentumCredit_keywords = r"\bmomentum\b"

    # Match categories by priority
    if re.search(fuel_keywords, text):
        return "Fuel"
    if re.search(shopping_keywords, text):
        return "Shopping"
    if re.search(utility_keywords, text):
        return "Utilities"
    if re.search(airtime_keywords, text):
        return "Airtime/Data"
    if re.search(betting_keywords, text):
        return "Betting"
    if re.search(paybill_keywords, text):
        return "Pay Bill"
    if re.search(buy_goods_keywords, text):
        return "Buy Goods"
    if re.search(agent_keywords, text):
        return "Agent Withdrawal"
    if re.search(salary_keywords, text):
        return "Income"
    if re.search(loan_keywords, text):
        return "Loan Repayment"
    if re.search(watucredit, text):
        return "WatuCredit"
    if re.search(momentum, text):
        return "MomentumCredit"
    if re.search(platinum, text):
        return "PlatinumCredit"
    if re.search(mogo, text):
        return "MogoCredit"
    else:
        return "Other"


def process_mpesa(text):
    lines = text.split('\n')
    transactions = []

    for i, line in enumerate(lines):
        if "Completed" in line:
            match = re.search(r'Completed[\s-]*(-?\d{1,3}(?:,\d{3})*(?:\.\d{2}))', line)
            if not match and i + 1 < len(lines):
                match = re.search(r'(-?\d{1,3}(?:,\d{3})*(?:\.\d{2}))', lines[i + 1])

            if match:
                amount_str = match.group(1).replace(",", "")
                try:
                    amount = float(amount_str)
                    Details = (line + " " + lines[i + 1]) if i + 1 < len(lines) else line
                    base_category = categorize_mpesa(Details)

                    #details_lines = [line]
                    #for j in range(1, 20):  # You can increase the range if needed
                    #    if i + j < len(lines):
                    #       details_lines.append(lines[i + j])
                    #Details = " ".join(details_lines)

                    details_lines = [line]
                    for j in range(1, len(lines) - i):
                        if "Completed" in lines[i + j]:  # Stop if next transaction starts
                            break
                        details_lines.append(lines[i + j])
                    Details = " ".join(details_lines)

                    base_category = categorize_mpesa(Details)

                    # Determine inflow/outflow
                    direction = "Inflow" if amount > 0 else "Outflow"

                    category = base_category

                    if category == "Other":
                        category = f"Other ({direction})"
                    # Specific overrides for loan disbursement vs repayment
                    #if base_category == "Loan":
                    #    category = "Loan Repayment" if amount > 0 else "Loan Disbursement"
                    #elif base_category == "Other":
                    #    category = "Other (Inflow)" if amount > 0 else "Other (Outflow)"
                    #else:
                    #    category = base_category

                    transactions.append({
                        "Details": Details.strip(),
                        "Amount": abs(amount),
                        "Inflow/Outflow": direction,
                        "Category": category
                    })
                except ValueError:
                    continue

    if not transactions:
        return None, None

    df = pd.DataFrame(transactions)

    # Group and summarize properly
    summary = df.groupby("Category")["Amount"].sum().reset_index()
    #summary = df.groupby("Category")["Amount"].sum(), count=("Amount", "count").reset_index()
    #summary = df.groupby("Category").size().reset_index(name="Count")
    #total_count = summary["count"].sum()
    total = df["Amount"].sum()
    #summary["Percentage"] = (summary["Amount"] / total * 100).round(2)
    summary["Count"] = df.groupby("Category")["Amount"].count().values

    return df, summary

#Bank Statements
def extract_bank(text):
    # Remove multiple blank lines to ensure consistency
    text = re.sub(r'\n\s*\n+', '\n', text)
    text = re.sub(r'[ \t]+', ' ', text)  # Replace tabs/multiple spaces with a single space
    text = text.strip()  # Remove leading/trailing whitespace

    # Extract account holder name
    name_match = re.search(r'Account Holder Name:\s*(.*)', text)
    account_holder_name = name_match.group(1).strip() if name_match else "N/A"

    # Extract account number
    account_number_match = re.search(r'Account Number:\s*(\d+)', text)
    account_number = account_number_match.group(1).strip() if account_number_match else "N/A"

    # Extract bank name
    bank_name_match = re.search(r'Bank Name:\s*(.*)', text)
    bank_name = bank_name_match.group(1).strip() if bank_name_match else "N/A"
    # Extract bank branch branch_name_match = re.search(r'Branch Name:\s*(.*)', text)
    if bank_file:
        bank_text = extract_text(bank_file, password=pdf_password)
    
    with st.expander("View Extracted Bank Text"):
        st.write(bank_text)

    bank_summary_info = extract_bank(bank_text)
    st.subheader("Bank Summary Info")
    st.json(bank_summary_info)

    bank_df, bank_summary = process_bank(bank_text)
    if bank_summary is not None:
        st.subheader("Bank Transaction Analysis")
        st.dataframe(bank_summary)
        st.bar_chart(bank_summary.set_index("Category")["Count"])
    else:
        st.warning("No valid bank transactions found.")

    return {
        "Account Holder Name": account_holder_name,
        "Account Number": account_number,
        "Bank Name": bank_name
    }

def process_bank(text):
    lines = text.split('\n')
    transactions = []
    
    for line in lines:
        line = line.strip()
        # Example pattern: "12/07/2025 POS Naivas -2,500.00"
        match = re.search(r'(\d{1,2}/\d{1,2}/\d{2,4})\s+(.+?)\s+(-?\d{1,3}(?:,\d{3})*(?:\.\d{2}))', line)
        if match:
            date_str, details, amount_str = match.groups()
            try:
                amount = float(amount_str.replace(",", ""))
                direction = "Inflow" if amount > 0 else "Outflow"
                category = categorize_mpesa(details)  # Reusing M-PESA logic

                if category == "Other":
                    category = f"Other ({direction})"

                transactions.append({
                    "Date": date_str,
                    "Details": details.strip(),
                    "Amount": abs(amount),
                    "Inflow/Outflow": direction,
                    "Category": category
                })
            except:
                continue

    if not transactions:
        return None, None

    df = pd.DataFrame(transactions)
    summary = df.groupby("Category")["Amount"].sum().reset_index()
    summary["Count"] = df.groupby("Category")["Amount"].count().values

    return df, summary

# CRB SECTION 
def extract_crb_data(text):
    # Remove multiple blank lines to ensure consistency
    text = re.sub(r'\n\s*\n+', '\n', text)
    text = re.sub(r'[ \t]+', ' ', text)  # Replace tabs/multiple spaces with a single space
    text = text.strip()  # Remove leading/trailing whitespace


    name_match = re.search(r'REPORTED NAMES:\s+(.*)', text)
    id_match = re.search(r'NATIONAL ID\s+:\s+(\d+)', text)
    #phone_match = re.search(r'Phone Number\(s\)[^\n]*\n\s*([\d, ]+)', text)
    #email_match = re.search(r'Email Address[^\n]*\n\s*([^\s]+@[^\s]+)', text)

    bio_data = {
        "Name": name_match.group(1).strip() if name_match else None,
        "National ID": id_match.group(1).strip() if id_match else "N/A",
        #"Phone Number(s)": phone_match.group(1).strip() if phone_match else "N/A",
        #"Email": email_match.group(1).strip() if email_match else "N/A",
    }
    pattern = r"\n\s*(\d+)\s*\n\s*(M\d)\s*\n\s*(\d+\s?%)"
    #pattern = r"Metro-Score©\s*(\d+)\s+PPI©\s*(M\d)\s+Probability Of Default©\s*(\d+\s?%)"
    #matches = re.findall(pattern, text, re.IGNORECASE | re.DOTALL)

    #pattern = r"Metro-Score©[\s\S]*?(\d+)[\s\S]*?PPI©[\s\S]*?(M\d)[\s\S]*?Probability Of Default©[\s\S]*?(\d+\s?%)"
    matches = re.findall(pattern, text, re.IGNORECASE| re.DOTALL)

    #for match in matches:
    #    print("Found block:", match)
    #breakpoint()

    if matches:
        metro, ppi, pd = matches[0]
    else:
        #metro = re.search(r'Metro-Score©\s+(\d+)', text)
        #ppi = re.search(r'PPI©\s+(M\d)', text)
        #pd = re.search(r'Probability Of Default©\s+(\d+\s?%)', text)

        #metro = int(metro.group(1)) if metro else None
        #ppi = ppi.group(1) if ppi else None
        #pd = pd.group(1).strip() if pd else None
        metro = re.search(r'Metro-Score©\s*\n*\s*(\d+)', text)
        ppi = re.search(r'PPI©\s*\n*\s*(M\d)', text)
        pd = re.search(r'Probability Of Default©\s*\n*\s*(\d+\s?%)', text)

        metro = metro.group(1) if metro else None
        ppi = ppi.group(1) if ppi else None
        pd = pd.group(1).strip() if pd else None

    credit_scores = {
        "Metro-Score": int(metro) if isinstance(metro, str) and metro.isdigit() else metro,
        "PPI": ppi,
        "Probability of Default": pd,
        "Interpretation": ""
    }
    #credit_scores = {
     #   "Metro-Score": int(metro.group(1)) if metro else "N/A",
      #  "PPI": ppi.group(1) if ppi else "N/A",
       # "Probability of Default": pd.group(1).strip() if pd else "N/A",
        #"Interpretation": ""
    #}

    metro_val = int(metro) if metro else None
    #ppi_val = ppi.group(1) if ppi else None
    ppi_val = ppi
    if metro_val:
        if metro_val < 400:
            credit_scores["Interpretation"] += "High Risk: Credit score indicates possible defaults.\n"
        elif metro_val < 600:
            credit_scores["Interpretation"] += "Medium Risk: Caution advised.\n"
        else:
            credit_scores["Interpretation"] += "Low Risk: Good credit standing.\n"
    if ppi_val:
        if ppi_val in ["M1", "M2"]:
            credit_scores["Interpretation"] += "Probable positive repayment behavior."
        elif ppi_val in ["M3", "M4", "M5"]:
            credit_scores["Interpretation"] += "Watch for occasional delays."
        else:
            credit_scores["Interpretation"] += "Probable Poor repayment trend."

    emp_match = re.search(r'Employer\s*:\s*(.*)', text, re.IGNORECASE)
    salary_match = re.search(r'Salary\s*:\s*K?([\d,]+)', text, re.IGNORECASE)
    dept_match = re.search(r'Department\s*:\s*(.*)', text, re.IGNORECASE)

    employment = {
        "Employer": emp_match.group(1).strip() if emp_match else "N/A",
        "Salary": salary_match.group(1).replace(",", "") if salary_match else "N/A",
        "Department": dept_match.group(1).strip() if dept_match else "N/A"
    }

    account_match = re.search(r'Total\s+(\d+)\s+(\d+)\s+(\d+)\s+(\d+)', text)
    balance_match = re.search(r'Total Outstanding Balance\s*\n\s*Total Accounts\s*\n\s*([\d,]+\.\d+)', text)
    balance_val = balance_match.group(1) if balance_match else "0.00"
    account_summary = {
        "Total Accounts": int(account_match.group(1)) if account_match else 0,
        "Non-Performing Accounts": int(account_match.group(2)) if account_match else 0,
        "Performing Accounts With Default History": int(account_match.group(3)) if account_match else 0,
        "Performing Accounts Without Default History": int(account_match.group(4)) if account_match else 0,
        "Total Outstanding Balance": float(balance_val.replace(",", "")) if balance_val else 0.0
    }

    return {
        "Bio Data": bio_data,
        "Employment": employment,
        "Credit Scores": credit_scores,
        "Account Summary": account_summary
    }

#def extract_crb_scores(text):
#    metro = re.search(r'Metro-Score©\s+(\d+)', text, re.IGNORECASE)
#    ppi = re.search(r'PPI©\s+([A-Z0-9]+)', text, re.IGNORECASE)
 #   pd = re.search(r'Probability Of Default©\s+(\d+ ?%)', text, re.IGNORECASE)

#    return {
#        "Metro Score": metro.group(1) if metro else "N/A",
#        "PPI": ppi.group(1) if ppi else "N/A",
#       "Probability of Default": pd.group(1) if pd else "N/A"
 #   }

# FILE UPLOAD + DISPLAY 
st.header("Upload M-PESA Statement (.txt, .pdf, .docx)")
mpesa_file = st.file_uploader("Upload M-PESA Statement", type=["txt", "pdf", "docx"], key="mpesa")
#st.button = st.button("Process M-PESA", key="process_mpesa")

st.header("Upload CRB Report (.txt, .pdf, .docx)")
crb_file = st.file_uploader("Upload CRB Report", type=["txt", "pdf", "docx"], key="crb")

st.header("Upload Bank Statement (.txt, .pdf, .docx)")
bank_file = st.file_uploader("Upload Bank Statement", type=["txt", "pdf", "docx"], key="bank")

st.info("If your file is encrypted, enter the password below.")
pdf_password = st.text_input("Enter PDF Password (optional):", type="password")


# Process M-PESA
if mpesa_file:
    mpesa_text = extract_text(mpesa_file, password=pdf_password)
    st.expander("View Extracted M-PESA Text").write(mpesa_text)  # Optional debug
    mpesa_df, mpesa_summary = process_mpesa(mpesa_text)

    if mpesa_summary is not None:
        st.subheader("M-PESA Expense Analysis")
        st.dataframe(mpesa_summary)
        st.bar_chart(mpesa_summary.set_index("Category")["Count"])
    else:
        st.warning("No valid M-PESA transactions found.")

# Process CRB
if crb_file:
    crb_text = extract_text(crb_file, password=pdf_password)
    crb_summary = extract_crb_data(crb_text)
    #crb_scores = extract_crb_scores(crb_text)
    #crb_scores = crb_summary["Credit Scores"]

    st.subheader("CRB Summary")
    st.json(crb_summary)

    #st.text_area("CRB Report Text", crb_text, height=300)
    #st.subheader("Credit Risk Scores")
    #st.json(crb_scores)

#if bank_file:
    bank_text = extract_text(bank_file, password=pdf_password) 
    bank_summary = extract_bank(bank_text)
    st.subheader("Bank Statement Summary")
    st.json(bank_summary)
